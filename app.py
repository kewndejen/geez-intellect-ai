import streamlit as st
import google.generativeai as genai
from PyPDF2 import PdfReader
from docx import Document
import time
import datetime

# ---------------------------------------------------------
# 1. IMPERIAL PAGE CONFIGURATION
# ---------------------------------------------------------
st.set_page_config(
    page_title="GE'EZ STUDIO | Sovereign Zenith",
    page_icon="🔱",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Majestic Emerald Green & Royal Gold Sovereign UI
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Cinzel+Decorative:wght@700;900&family=Montserrat:wght@400;700&family=Abyssinica+SIL&display=swap');
    
    .stApp { 
        background: radial-gradient(circle at center, #003311 0%, #001a0d 100%); 
        color: #ffffff; 
        font-family: 'Montserrat', sans-serif; 
    }
    
    h1, h2, h3 { 
        font-family: 'Cinzel Decorative', serif; 
        color: #FFD700 !important; 
        text-align: center;
        text-shadow: 0px 0px 20px rgba(255, 215, 0, 0.5);
    }

    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, #001a0d 0%, #000a05 100%) !important;
        border-right: 3px solid #FFD700;
    }
    
    .sovereign-card {
        background: rgba(255, 255, 255, 0.12);
        backdrop-filter: blur(15px);
        padding: 30px; border-radius: 20px;
        border: 1px solid rgba(255, 215, 0, 0.4);
        border-left: 15px solid #FFD700;
        margin-bottom: 25px;
    }

    .stButton>button {
        background: linear-gradient(135deg, #FFD700 0%, #B8860B 100%) !important;
        color: #000000 !important; font-weight: 900 !important;
        border-radius: 12px !important; border: 2px solid #FFFFFF !important;
        height: 3.8em; width: 100%; transition: 0.5s ease;
        text-transform: uppercase;
    }
    .stButton>button:hover { transform: translateY(-3px); box-shadow: 0 0 35px #FFD700; }

    [data-testid="stChatInput"] { 
        border: 3px solid #FFD700 !important; 
        background-color: #ffffff !important; 
        border-radius: 20px !important; 
    }
    [data-testid="stChatInput"] textarea { color: #000000 !important; font-weight: bold !important; font-size: 1.1rem !important; }
    
    .wait-msg { color: #FFD700; font-style: italic; font-size: 1rem; text-align: center; font-weight: bold; background: rgba(0,0,0,0.3); padding: 10px; border-radius: 10px; }
    .citation { font-size: 0.85rem; color: #FFD700; border-top: 1px solid rgba(255,255,255,0.2); margin-top: 25px; padding-top: 15px; font-family: monospace; }
    </style>
    """, unsafe_allow_html=True)

# ---------------------------------------------------------
# 2. DOCUMENT ARCHIVE ENGINE (PDF/DOCX)
# ---------------------------------------------------------
def extract_text(uploaded_file):
    ext = uploaded_file.name.split('.')[-1].lower()
    text = ""
    try:
        if ext == 'pdf':
            reader = PdfReader(uploaded_file)
            for page in reader.pages: text += page.extract_text()
        elif ext in ['docx', 'doc']:
            doc = Document(uploaded_file)
            for p in doc.paragraphs: text += p.text + "\n"
        return text
    except Exception as e:
        return f"Error reading file: {e}"

# ---------------------------------------------------------
# 3. INDESTRUCTIBLE AI ENGINE (Multi-Model Support)
# ---------------------------------------------------------
if "GOOGLE_API_KEY" in st.secrets:
    genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])
else:
    st.error("⚠️ API Key አልተገኘም!")
    st.stop()

@st.cache_resource
def get_stable_engine():
    """የእርስዎ ቁልፍ የሚፈቅዳቸውን ሞዴሎች በራስ-ሰር ይለያል"""
    try:
        available = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        # 1.5-flash እጅግ ሰፊ ኮታ ስላለው ለሥራ አይቆራረጥም
        priority = ["models/gemini-1.5-flash", "models/gemini-1.5-pro", "models/gemini-pro"]
        for target in priority:
            for actual in available:
                if target in actual: return actual
        return available[0]
    except:
        return "models/gemini-1.5-flash"

ACTIVE_ENGINE = get_stable_engine()

def ask_sovereign_scholar(prompt, tool_context, document_archive=""):
    sys_instr = f"""
    You are 'Ge'ez Scholar AI Master', a Gemini 3 standard intelligence created by Grand Architect Deacon Kewn Dejen.
    Expertise Area: {tool_context}. Knowledge: 60 Pillars of Wisdom.
    
    PRIMARY SOURCE (DOCUMENT VAULT):
    {document_archive[:25000]}
    
    Task: Provide scholarly, historical, and deep analysis. Support Ge'ez/Amharic.
    Tone: Sovereign, authoritative, and extremely wise.
    """
    
    status_placeholder = st.empty()
    
    # ተስፋ የማይቆርጥ ሎጂክ (Retry loop to handle 429)
    for attempt in range(1, 4):
        try:
            model = genai.GenerativeModel(model_name=ACTIVE_ENGINE, system_instruction=sys_instr)
            response = model.generate_content(prompt)
            if response and response.text:
                status_placeholder.empty()
                return response.text, ACTIVE_ENGINE
        except Exception as e:
            if "429" in str(e):
                wait = (attempt * 7)
                status_placeholder.markdown(f"<div class='wait-msg'>⏳ ሊቁ በጥልቅ ምርምር ላይ ናቸው... (ሙከራ {attempt}/3)</div>", unsafe_allow_html=True)
                time.sleep(wait)
                continue
            break
            
    status_placeholder.empty()
    return "❌ ሊቁ በአሁኑ ሰዓት መዛግብቱን ለመክፈት አልቻሉም (ኮታ አልቋል)። እባክዎ ከ1 ደቂቃ በኋላ ገጹን Refresh አድርገው ይሞክሩ።", "None"

# ---------------------------------------------------------
# 4. SIDEBAR: THE ARK OF 60 PILLARS
# ---------------------------------------------------------
with st.sidebar:
    st.markdown("<h1>🔱 GE'EZ STUDIO</h1>", unsafe_allow_html=True)
    st.markdown(f"""
        <div style='background:linear-gradient(45deg, #FFD700, #B8860B); padding:15px; border-radius:12px; text-align:center; color:#000; font-weight:900; border: 2px solid #fff;'>
            GRAND ARCHITECT:<br>DEACON KEWN DEJEN
        </div>
    """, unsafe_allow_html=True)
    st.markdown("---")

    # 📂 Gemini 3 Multi-Doc Vault
    st.subheader("📂 የሰነድ መዝገብ (Document Vault)")
    uploaded_files = st.file_uploader("ሰነዶችን እዚህ ይደረድሩ (PDF/Word)", type=['pdf', 'docx'], accept_multiple_files=True)
    
    if "global_memory" not in st.session_state: st.session_state.global_memory = ""
    if "file_names" not in st.session_state: st.session_state.file_names = []

    if uploaded_files:
        combined_text = ""
        current_names = [f.name for f in uploaded_files]
        if current_names != st.session_state.file_names:
            with st.spinner("ጀሚኒ 3 ሰነዶቹን እየተነተነ ነው..."):
                for f in uploaded_files:
                    combined_text += f"\n[FILE: {f.name}]\n" + extract_text(f)
                st.session_state.global_memory = combined_text
                st.session_state.file_names = current_names
                st.success(f"✅ {len(uploaded_files)} ሰነዶች ተነበዋል!")

    st.markdown("---")
    pillar = st.selectbox("የጥበብ ምሰሶ ይምረጡ", [
        "Advanced AI Labs", "Digital Archives", "Heritage & Science",
        "Imperial University", "Mysticism & Qene", "Strategic Wealth"
    ])

    # 60 Tools Detailed Map
    tools_map = {
        "Advanced AI Labs": ["Manuscript OCR", "Linguistic Bridge", "Script Authentication", "Root Finder", "Voice of Wisdom", "Neural Translation", "Syntax Analyzer", "Ge'ez NLP", "Dialect Study", "Semantic Map"],
        "Digital Archives": ["Universal Library", "Fetha Nagast AI", "Synaxarium Analysis", "Royal Decrees", "Treaty Expert", "Kibre Nagast Hub", "Ecclesiastical Law", "Genealogy Map", "Preservation Lab", "Hagiography Lab"],
        "Heritage & Science": ["Ancient Medicine", "Archeology Hub", "Heritage Map", "Iconography Vision", "Architecture AI", "Geology of Axum", "Botany Hub", "Zoology Hub", "Ink Chemistry", "Virtual Museum"],
        "Imperial University": ["Bahre Hasab Pro", "Abu Shaker Astronomy", "Numerology Lab", "Scribe Assistant", "Font Converter", "Ethiopic Math", "Philosophy Hub", "History Chronology", "University Portal", "Scholarly Citation"],
        "Mysticism & Qene": ["Sem-na-Worq (Qene)", "St. Yared Zema Lab", "Theology Hub", "Scholar Roleplay", "Proverbs AI", "Esoteric Wisdom", "Liturgical Guide", "Monastic Study", "Apostolic Tradition", "Hymnology Lab"],
        "Strategic Wealth": ["Business Hub", "API Portal", "Security Admin", "Wealth Strategy", "Sovereign Logs", "Global Relations", "Grant Assistant", "Strategic Planning", "Project Manager", "Data Vault"]
    }
    tool = st.radio("Labs", tools_map[pillar])

    if st.button("🔄 REBOOT SYSTEM"):
        st.session_state.global_memory = ""
        st.session_state.file_names = []
        st.session_state.messages = []
        st.cache_resource.clear()
        st.rerun()

# ---------------------------------------------------------
# 5. MAIN WORKSPACE
# ---------------------------------------------------------
st.markdown(f"<h1>{tool}</h1>", unsafe_allow_html=True)

if st.session_state.file_names:
    with st.expander("📝 በመዝገብ ላይ ያሉ ሰነዶች"):
        for name in st.session_state.file_names:
            st.write(f"🔹 {name}")

if "messages" not in st.session_state: st.session_state.messages = []
for m in st.session_state.messages:
    with st.chat_message(m["role"]):
        st.markdown(m["content"], unsafe_allow_html=True)

if prompt := st.chat_input("ለሊቁ ጥያቄዎን እዚህ ያስገቡ..."):
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(f"<div style='color: white;'><b>{prompt}</b></div>", unsafe_allow_html=True)

    with st.chat_message("assistant"):
        with st.spinner("ሊቁ ሰነዶቹንና መዛግብቱን በጥልቀት እያመሳከረ ነው..."):
            answer, engine = ask_sovereign_scholar(prompt, tool, st.session_state.global_memory)
            
            full_res = f"""
            <div style='color: white; line-height: 1.8;'>{answer}</div>
            <div class='citation'>Source: {engine} | v-Masterpiece Sovereign Edition</div>
            """
            st.markdown(full_res, unsafe_allow_html=True)
            st.session_state.messages.append({"role": "assistant", "content": full_res})

# Master Footer
st.markdown("<br><hr><p style='text-align:center; color:#FFD700;'><b>GE'EZ SCHOLAR AI STUDIO | Grand Architect Deacon Kewn Dejen</b><br>© 2024-2025 ALL RIGHTS RESERVED | THE MASTERPIECE EDITION</p>", unsafe_allow_html=True)import streamlit as st
import google.generativeai as genai
from PyPDF2 import PdfReader
from docx import Document
import time
import datetime

# ---------------------------------------------------------
# 1. IMPERIAL PAGE CONFIGURATION
# ---------------------------------------------------------
st.set_page_config(
    page_title="GE'EZ STUDIO | Sovereign Zenith",
    page_icon="🔱",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Majestic Emerald Green & Royal Gold Sovereign UI
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Cinzel+Decorative:wght@700;900&family=Montserrat:wght@400;700&family=Abyssinica+SIL&display=swap');
    
    .stApp { 
        background: radial-gradient(circle at center, #003311 0%, #001a0d 100%); 
        color: #ffffff; 
        font-family: 'Montserrat', sans-serif; 
    }
    
    h1, h2, h3 { 
        font-family: 'Cinzel Decorative', serif; 
        color: #FFD700 !important; 
        text-align: center;
        text-shadow: 0px 0px 20px rgba(255, 215, 0, 0.5);
    }

    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, #001a0d 0%, #000a05 100%) !important;
        border-right: 3px solid #FFD700;
    }
    
    .sovereign-card {
        background: rgba(255, 255, 255, 0.12);
        backdrop-filter: blur(15px);
        padding: 30px; border-radius: 20px;
        border: 1px solid rgba(255, 215, 0, 0.4);
        border-left: 15px solid #FFD700;
        margin-bottom: 25px;
    }

    .stButton>button {
        background: linear-gradient(135deg, #FFD700 0%, #B8860B 100%) !important;
        color: #000000 !important; font-weight: 900 !important;
        border-radius: 12px !important; border: 2px solid #FFFFFF !important;
        height: 3.8em; width: 100%; transition: 0.5s ease;
        text-transform: uppercase;
    }
    .stButton>button:hover { transform: translateY(-3px); box-shadow: 0 0 35px #FFD700; }

    [data-testid="stChatInput"] { 
        border: 3px solid #FFD700 !important; 
        background-color: #ffffff !important; 
        border-radius: 20px !important; 
    }
    [data-testid="stChatInput"] textarea { color: #000000 !important; font-weight: bold !important; font-size: 1.1rem !important; }
    
    .wait-msg { color: #FFD700; font-style: italic; font-size: 1rem; text-align: center; font-weight: bold; background: rgba(0,0,0,0.3); padding: 10px; border-radius: 10px; }
    .citation { font-size: 0.85rem; color: #FFD700; border-top: 1px solid rgba(255,255,255,0.2); margin-top: 25px; padding-top: 15px; font-family: monospace; }
    </style>
    """, unsafe_allow_html=True)

# ---------------------------------------------------------
# 2. DOCUMENT ARCHIVE ENGINE (PDF/DOCX)
# ---------------------------------------------------------
def extract_text(uploaded_file):
    ext = uploaded_file.name.split('.')[-1].lower()
    text = ""
    try:
        if ext == 'pdf':
            reader = PdfReader(uploaded_file)
            for page in reader.pages: text += page.extract_text()
        elif ext in ['docx', 'doc']:
            doc = Document(uploaded_file)
            for p in doc.paragraphs: text += p.text + "\n"
        return text
    except Exception as e:
        return f"Error reading file: {e}"

# ---------------------------------------------------------
# 3. INDESTRUCTIBLE AI ENGINE (Multi-Model Support)
# ---------------------------------------------------------
if "GOOGLE_API_KEY" in st.secrets:
    genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])
else:
    st.error("⚠️ API Key አልተገኘም!")
    st.stop()

@st.cache_resource
def get_stable_engine():
    """የእርስዎ ቁልፍ የሚፈቅዳቸውን ሞዴሎች በራስ-ሰር ይለያል"""
    try:
        available = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        # 1.5-flash እጅግ ሰፊ ኮታ ስላለው ለሥራ አይቆራረጥም
        priority = ["models/gemini-1.5-flash", "models/gemini-1.5-pro", "models/gemini-pro"]
        for target in priority:
            for actual in available:
                if target in actual: return actual
        return available[0]
    except:
        return "models/gemini-1.5-flash"

ACTIVE_ENGINE = get_stable_engine()

def ask_sovereign_scholar(prompt, tool_context, document_archive=""):
    sys_instr = f"""
    You are 'Ge'ez Scholar AI Master', a Gemini 3 standard intelligence created by Grand Architect Deacon Kewn Dejen.
    Expertise Area: {tool_context}. Knowledge: 60 Pillars of Wisdom.
    
    PRIMARY SOURCE (DOCUMENT VAULT):
    {document_archive[:25000]}
    
    Task: Provide scholarly, historical, and deep analysis. Support Ge'ez/Amharic.
    Tone: Sovereign, authoritative, and extremely wise.
    """
    
    status_placeholder = st.empty()
    
    # ተስፋ የማይቆርጥ ሎጂክ (Retry loop to handle 429)
    for attempt in range(1, 4):
        try:
            model = genai.GenerativeModel(model_name=ACTIVE_ENGINE, system_instruction=sys_instr)
            response = model.generate_content(prompt)
            if response and response.text:
                status_placeholder.empty()
                return response.text, ACTIVE_ENGINE
        except Exception as e:
            if "429" in str(e):
                wait = (attempt * 7)
                status_placeholder.markdown(f"<div class='wait-msg'>⏳ ሊቁ በጥልቅ ምርምር ላይ ናቸው... (ሙከራ {attempt}/3)</div>", unsafe_allow_html=True)
                time.sleep(wait)
                continue
            break
            
    status_placeholder.empty()
    return "❌ ሊቁ በአሁኑ ሰዓት መዛግብቱን ለመክፈት አልቻሉም (ኮታ አልቋል)። እባክዎ ከ1 ደቂቃ በኋላ ገጹን Refresh አድርገው ይሞክሩ።", "None"

# ---------------------------------------------------------
# 4. SIDEBAR: THE ARK OF 60 PILLARS
# ---------------------------------------------------------
with st.sidebar:
    st.markdown("<h1>🔱 GE'EZ STUDIO</h1>", unsafe_allow_html=True)
    st.markdown(f"""
        <div style='background:linear-gradient(45deg, #FFD700, #B8860B); padding:15px; border-radius:12px; text-align:center; color:#000; font-weight:900; border: 2px solid #fff;'>
            GRAND ARCHITECT:<br>DEACON KEWN DEJEN
        </div>
    """, unsafe_allow_html=True)
    st.markdown("---")

    # 📂 Gemini 3 Multi-Doc Vault
    st.subheader("📂 የሰነድ መዝገብ (Document Vault)")
    uploaded_files = st.file_uploader("ሰነዶችን እዚህ ይደረድሩ (PDF/Word)", type=['pdf', 'docx'], accept_multiple_files=True)
    
    if "global_memory" not in st.session_state: st.session_state.global_memory = ""
    if "file_names" not in st.session_state: st.session_state.file_names = []

    if uploaded_files:
        combined_text = ""
        current_names = [f.name for f in uploaded_files]
        if current_names != st.session_state.file_names:
            with st.spinner("ጀሚኒ 3 ሰነዶቹን እየተነተነ ነው..."):
                for f in uploaded_files:
                    combined_text += f"\n[FILE: {f.name}]\n" + extract_text(f)
                st.session_state.global_memory = combined_text
                st.session_state.file_names = current_names
                st.success(f"✅ {len(uploaded_files)} ሰነዶች ተነበዋል!")

    st.markdown("---")
    pillar = st.selectbox("የጥበብ ምሰሶ ይምረጡ", [
        "Advanced AI Labs", "Digital Archives", "Heritage & Science",
        "Imperial University", "Mysticism & Qene", "Strategic Wealth"
    ])

    # 60 Tools Detailed Map
    tools_map = {
        "Advanced AI Labs": ["Manuscript OCR", "Linguistic Bridge", "Script Authentication", "Root Finder", "Voice of Wisdom", "Neural Translation", "Syntax Analyzer", "Ge'ez NLP", "Dialect Study", "Semantic Map"],
        "Digital Archives": ["Universal Library", "Fetha Nagast AI", "Synaxarium Analysis", "Royal Decrees", "Treaty Expert", "Kibre Nagast Hub", "Ecclesiastical Law", "Genealogy Map", "Preservation Lab", "Hagiography Lab"],
        "Heritage & Science": ["Ancient Medicine", "Archeology Hub", "Heritage Map", "Iconography Vision", "Architecture AI", "Geology of Axum", "Botany Hub", "Zoology Hub", "Ink Chemistry", "Virtual Museum"],
        "Imperial University": ["Bahre Hasab Pro", "Abu Shaker Astronomy", "Numerology Lab", "Scribe Assistant", "Font Converter", "Ethiopic Math", "Philosophy Hub", "History Chronology", "University Portal", "Scholarly Citation"],
        "Mysticism & Qene": ["Sem-na-Worq (Qene)", "St. Yared Zema Lab", "Theology Hub", "Scholar Roleplay", "Proverbs AI", "Esoteric Wisdom", "Liturgical Guide", "Monastic Study", "Apostolic Tradition", "Hymnology Lab"],
        "Strategic Wealth": ["Business Hub", "API Portal", "Security Admin", "Wealth Strategy", "Sovereign Logs", "Global Relations", "Grant Assistant", "Strategic Planning", "Project Manager", "Data Vault"]
    }
    tool = st.radio("Labs", tools_map[pillar])

    if st.button("🔄 REBOOT SYSTEM"):
        st.session_state.global_memory = ""
        st.session_state.file_names = []
        st.session_state.messages = []
        st.cache_resource.clear()
        st.rerun()

# ---------------------------------------------------------
# 5. MAIN WORKSPACE
# ---------------------------------------------------------
st.markdown(f"<h1>{tool}</h1>", unsafe_allow_html=True)

if st.session_state.file_names:
    with st.expander("📝 በመዝገብ ላይ ያሉ ሰነዶች"):
        for name in st.session_state.file_names:
            st.write(f"🔹 {name}")

if "messages" not in st.session_state: st.session_state.messages = []
for m in st.session_state.messages:
    with st.chat_message(m["role"]):
        st.markdown(m["content"], unsafe_allow_html=True)

if prompt := st.chat_input("ለሊቁ ጥያቄዎን እዚህ ያስገቡ..."):
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(f"<div style='color: white;'><b>{prompt}</b></div>", unsafe_allow_html=True)

    with st.chat_message("assistant"):
        with st.spinner("ሊቁ ሰነዶቹንና መዛግብቱን በጥልቀት እያመሳከረ ነው..."):
            answer, engine = ask_sovereign_scholar(prompt, tool, st.session_state.global_memory)
            
            full_res = f"""
            <div style='color: white; line-height: 1.8;'>{answer}</div>
            <div class='citation'>Source: {engine} | v-Masterpiece Sovereign Edition</div>
            """
            st.markdown(full_res, unsafe_allow_html=True)
            st.session_state.messages.append({"role": "assistant", "content": full_res})

# Master Footer
st.markdown("<br><hr><p style='text-align:center; color:#FFD700;'><b>GE'EZ SCHOLAR AI STUDIO | Grand Architect Deacon Kewn Dejen</b><br>© 2024-2025 ALL RIGHTS RESERVED | THE MASTERPIECE EDITION</p>", unsafe_allow_html=True)
